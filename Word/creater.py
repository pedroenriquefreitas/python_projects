import docx

d = docx.Document('/Users/pedroenriqueandrade/Desktop/python_projects/Word/text.docx')

tables = d.tables

for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if 'Entrevistas' in paragraph.text:
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None

for paragraph in d.paragraphs:
    if 'Entrevistas' in paragraph.text:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

d.save('/Users/pedroenriqueandrade/Desktop/python_projects/Word/text.docx')
